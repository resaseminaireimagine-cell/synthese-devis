# app.py — Synthèse devis prestataires (FINAL+++)
# ---------------------------------------------
# Fixes:
# - Supprime la page blanche après "DÉTAIL DES OFFRES..." (plus de page_break juste après)
# - Vendor béton: interdit "SAS"/"SARL" seuls, découpe texte collé, coupe avant adresse
# - Filtrage logistique/admin renforcé (adresses / contacts / headers / IBAN / etc.)
# - Tech: TVA supprimée sans créer "TV" ; blacklist tabular label "TV/TVA"
# - UX: onglets "Synthèse" vs "Détail" conservés

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional

import streamlit as st
from pypdf import PdfReader

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# BRAND
# =========================
APP_TITLE = "Synthèse devis prestataires — Institut Imagine"
PRIMARY = "#AF0073"
BG = "#F6F7FB"
FONT = "Montserrat"


CATERING_POSTS = [
    "Accueil café",
    "Pause matin",
    "Déjeuner",
    "Pause après-midi",
    "Cocktail",
    "Boissons (global)",
    "Options",
    "Autres (logistique)",
]

TECH_POSTS = [
    "Périmètre",
    "Équipe",
    "Captation",
    "Régie",
    "Diffusion",
    "Replay",
    "Inclus",
    "Contraintes / options",
    "Conseil",
]


# =========================
# LEXICON
# =========================
MENU_HINTS = [
    "accueil", "petit", "déjeuner", "dejeuner", "pause", "buffet", "cocktail", "apéritif", "aperitif",
    "déjeunatoire", "dejeunatoire",
    "café", "cafe", "thé", "the", "soft", "jus", "eau",
    "viennoiser", "gourmand", "mignard", "financier", "cannel",
    "pièce", "pieces", "pièces", "/pers", "par personne", "convive", "invité", "invite",
    "salée", "sucrée", "sucree",
    "sandwich", "wrap", "salade", "fromage", "fruit",
    "vin", "champagne", "bière", "biere",
    "thermos", "gobelet", "tasse", "serviette", "plateau",
    "verrerie", "flûtes", "assiettes", "nappage", "mobilier", "mange-debout",
    "livraison", "reprise", "mise en place", "personnel", "service",
]

TECH_HINTS = [
    "captation", "caméra", "camera", "4k", "cadreur", "réalisateur", "realisateur",
    "ingénieur", "ingenieur", "son", "audio",
    "régie", "regie", "diffusion", "live", "zoom", "duplex", "plateforme",
    "replay", "wetransfer", "we transfer", "enregistrement",
    "pavlov", "zapette", "tv", "écran", "ecran",
    "micro", "hf", "console", "mélangeur", "melangeur", "obs", "vmix",
]

LEGAL_FORMS = ["sas", "sarl", "sa", "eurl", "sasu", "association", "scop", "groupe"]

ADDRESS_HINTS = [
    "rue", "avenue", "boulevard", "allée", "allee", "bp", "cedex",
    "france", "paris", "clichy", "nanterre", "saint",
    "quai", "impasse", "route", "chemin",
]

VENDOR_FORBIDDEN = [
    "en euros", "appliqué", "applique", "devis", "facture",
    "désignation", "designation", "quantité", "quantite",
    "montant", "tva", "base ht", "total", "ttc", "ht",
    "récapitulatif", "recapitulatif", "proposition", "prestation",
    "prix par convive",
    "au capital",
]

ADMIN_HINTS = [
    "conditions générales", "conditions generales", "cgv",
    "rgpd", "données personnelles", "donnees personnelles",
    "iban", "bic", "rib", "banque", "tva intracommunautaire",
    "siret", "rcs", "capital",
    "adresse", "tél", "tel", "email", "e-mail", "site internet", "www.",
    "référence", "reference", "date de devis", "date de validité", "signature",
    "mode de paiement", "net à payer", "net a payer",
    "tribunal", "mise en demeure", "penalite", "pénalité",
    "indemnité forfaitaire", "indemnite forfaitaire",
    "page ",
    "code description", "montant ht", "montant tva", "base ht",
]


# =========================
# TEXT UTILS
# =========================
def norm(s: str) -> str:
    s = "" if s is None else str(s)
    s = s.replace("\u00A0", " ").replace("\t", " ").replace("\r", " ")
    s = s.strip()
    s = re.sub(r"\s+", " ", s)
    return s


def fold(s: str) -> str:
    return norm(s).lower()


def hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.replace("#", "").strip()
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def extract_pdf_text(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    return "\n".join([(p.extract_text() or "") for p in reader.pages])


def split_lines(text: str) -> List[str]:
    # bullets + decollage lettres/chiffres (ex: "Réceptions1" -> "Réceptions 1")
    text = text.replace("•", "\n• ")
    text = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ])(\d)", r"\1 \2", text)
    text = re.sub(r"(\d)([A-Za-zÀ-ÖØ-öø-ÿ])", r"\1 \2", text)

    raw = text.splitlines()
    out = []
    for r in raw:
        rr = r.replace("\u00A0", " ")
        rr = re.sub(r"\s+", " ", rr).strip()
        if rr:
            out.append(rr)
    return out


def cut_at_cgv(lines: List[str]) -> List[str]:
    out = []
    for ln in lines:
        l = fold(ln)
        if "conditions générales" in l or "conditions generales" in l or re.search(r"\bcgv\b", l):
            break
        out.append(ln)
    return out


def looks_like_schedule_line(s: str) -> bool:
    l = fold(s)
    return bool(re.search(r"\b0?\d{1,2}h\d{2}\b", l) and ((" à " in l) or (" a " in l) or ("-" in l)) and len(l) <= 220)


def looks_like_price_table_line(s: str) -> bool:
    core = re.sub(r"[€]", "", norm(s))
    digits = sum(ch.isdigit() for ch in core)
    letters = sum(ch.isalpha() for ch in core)
    if digits >= 10 and digits > letters:
        return True
    if re.fullmatch(r"[\d\s,\.%\-\/]+", core) and digits >= 6:
        return True
    return False


def is_addressy(s: str) -> bool:
    l = fold(s)
    has_cp = bool(re.search(r"\b\d{5}\b", l))
    strong = bool(re.search(r"^\s*\d{1,4}\s*(bis|ter)?\s*(rue|avenue|boulevard|quai|impasse|route|chemin)\b", l))
    hits = sum(k in l for k in ADDRESS_HINTS)
    return strong or (has_cp and hits >= 1) or (hits >= 2)


def is_admin_line(s: str) -> bool:
    l = fold(s)
    if any(k in l for k in ADMIN_HINTS):
        return True
    if re.search(r"\bpage\s+\d+\s+sur\s+\d+\b", l):
        return True
    compact = re.sub(r"\s+", "", s)
    if re.search(r"\bFR\d{2}[0-9A-Z]{10,30}\b", compact):
        return True
    if "@" in s:
        return True
    if re.search(r"\b0[1-9](\s?\d{2}){4}\b", s):
        return True
    return False


def is_noise_line(s: str) -> bool:
    s = norm(s)
    if not s:
        return True
    if looks_like_schedule_line(s):
        return False  # on garde les horaires dans le détail si utiles
    if looks_like_price_table_line(s):
        return False  # on gère via labels tabulaires
    if re.fullmatch(r"\d{1,6}", s):
        return True
    if is_admin_line(s):
        return True
    return False


def parse_eur_amount(s: str) -> Optional[float]:
    s = norm(s)
    if not s:
        return None
    s = s.replace("€", "").replace("EUR", "").replace("euros", "")
    s = re.sub(r"[^0-9,.\s-]", "", s).strip()
    if not s:
        return None
    s2 = s.replace(" ", "")
    if "," in s2 and "." in s2:
        s2 = s2.replace(".", "")
    s2 = s2.replace(",", ".")
    try:
        return float(s2)
    except Exception:
        return None


def euro_fmt(x: Optional[float]) -> str:
    if x is None:
        return "—"
    return f"{x:,.2f} €".replace(",", " ").replace(".", ",")


def find_total_ttc(text: str) -> Optional[float]:
    patterns = [
        r"total\s+ttc\s*[:\-]?\s*([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
        r"net\s+à\s+payer.*?([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
        r"net\s+a\s+payer.*?([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
        r"total\s+à\s+payer.*?([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
        r"total\s+devis\s+t\.t\.c\.\s*[:\-]?\s*([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
        r"total\s+g[ée]n[ée]ral\s+ttc.*?([0-9][0-9 \u00A0]*[,\.][0-9]{2})",
    ]
    lt = fold(text)
    found: List[float] = []
    for pat in patterns:
        for m in re.finditer(pat, lt, flags=re.IGNORECASE | re.DOTALL):
            amt = parse_eur_amount(m.group(1))
            if amt is not None:
                found.append(amt)
    return found[-1] if found else None


# =========================
# VENDOR
# =========================
def vendor_from_filename(filename: str) -> str:
    base = filename.rsplit(".", 1)[0]
    base = re.sub(r"[_\-]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip()
    base = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ])(\d)", r"\1 \2", base)

    junk = {"v1", "v2", "v3", "devis", "dev", "facture", "institut", "imagine", "pax", "100pax"}
    tokens = [t for t in re.split(r"\s+", base) if t]
    kept = []
    for t in tokens:
        tl = t.lower()
        if tl in junk:
            continue
        if re.fullmatch(r"\d{3,}", tl):
            continue
        if re.fullmatch(r"de[\-_]?\d+", tl):
            continue
        kept.append(t)

    out = []
    for t in kept:
        if re.search(r"\d", t) and len(out) >= 2:
            break
        out.append(t)
        if len(out) >= 5:
            break

    candidate = " ".join(out).strip()
    return candidate if candidate else base


def vendor_is_suspicious(v: str) -> bool:
    v = norm(v)
    lv = fold(v)
    if not v or len(v) < 3:
        return True
    # "SAS" / "SARL" seuls => interdit
    if lv in set(LEGAL_FORMS) or re.fullmatch(r"(sas|sarl|sa|eurl|sasu)\b\.?", lv):
        return True
    if any(k in lv for k in VENDOR_FORBIDDEN):
        return True
    if is_addressy(v):
        return True
    if "@" in v or "contact" in lv:
        return True
    if sum(ch.isdigit() for ch in v) >= 8 and sum(ch.isalpha() for ch in v) <= 4:
        return True
    return False


def normalize_company_line(line: str) -> str:
    s = norm(line)
    s = re.sub(r"([A-Za-zÀ-ÖØ-öø-ÿ])(\d)", r"\1 \2", s)
    l = fold(s)

    # coupe "au capital"
    if "au capital" in l:
        s = re.split(r"(?i)\bau capital\b", s)[0].strip(" -–,;:")

    # coupe avant adresse si " - 10 bis rue ..."
    if " - " in s and is_addressy(s.split(" - ", 1)[1]):
        s = s.split(" - ", 1)[0].strip()

    # coupe si l'adresse suit collée sans tiret ("Réceptions 1 rue ...")
    if re.search(r"\b(rue|avenue|boulevard|quai|impasse|route|chemin)\b", fold(s)):
        m = re.search(r"\b(rue|avenue|boulevard|quai|impasse|route|chemin)\b", fold(s))
        if m:
            left = s[:m.start()].strip(" -–,;:")
            if left and len(left) >= 3:
                s = left

    return s


def guess_vendor_name(text: str, filename: str) -> str:
    file_vendor = vendor_from_filename(filename)

    lines = [normalize_company_line(x) for x in text.splitlines() if norm(x)]
    top = lines[:260]

    # (1) ligne avec forme légale + un vrai nom derrière
    legal_candidates = []
    for ln in top:
        l = fold(ln)
        if any(f" {lf} " in f" {l} " for lf in LEGAL_FORMS):
            cand = normalize_company_line(ln)
            if not vendor_is_suspicious(cand) and sum(ch.isalpha() for ch in cand) >= 6:
                legal_candidates.append(cand)
    if legal_candidates:
        legal_candidates.sort(key=lambda x: len(x))
        return legal_candidates[0]

    # (2) au-dessus de SIRET
    for idx, ln in enumerate(top):
        l = fold(ln)
        if "siret" in l or "rcs" in l:
            for back in range(1, 14):
                j = idx - back
                if j >= 0:
                    cand = normalize_company_line(top[j])
                    if cand and not vendor_is_suspicious(cand) and sum(ch.isalpha() for ch in cand) >= 6:
                        return cand
            break

    # (3) fallback: filename
    return file_vendor


# =========================
# TABULAR label extraction
# =========================
def extract_left_label_from_tabular(line: str) -> Optional[str]:
    s = norm(line)
    if len(s) < 6:
        return None
    if is_admin_line(s):
        return None

    parts = s.split()
    cut = None
    for i, tok in enumerate(parts):
        if any(ch.isdigit() for ch in tok):
            cut = i
            break
    if cut is None or cut == 0:
        return None

    left = " ".join(parts[:cut]).strip(" -–:;")
    if len(left) < 4:
        return None

    ll = fold(left)
    # blacklist TVA/TV (sinon on récupère "TV" depuis "TVA :")
    if ll in {"tv", "tva", "total tva"}:
        return None
    if any(k in ll for k in ["total", "montant", "remise", "désignation", "designation", "quantité", "quantite", "base ht"]):
        return None
    if is_addressy(left):
        return None
    return left


# =========================
# DETAIL FILTERS
# =========================
def should_keep_catering_detail(item: str) -> bool:
    s = norm(item)
    l = fold(s)
    if not s:
        return False
    if is_admin_line(s):
        return False
    if is_addressy(s):
        return False

    # tue contacts / destinataires / références internes
    if any(k in l for k in ["institut imagine", "a l'attention", "votre contact", "contact client", "réf. interne", "ref. interne"]):
        return False
    if any(k in l for k in ["date de la prestation", "quai de livraison", "code description", "conserver au frais"]):
        return False
    if re.search(r"\b(désignation|designation|quantité|quantite|montant|p\.u|pu ht|base ht|total ht|total tva)\b", l):
        return False

    return True


def should_keep_tech_detail(item: str) -> bool:
    s = norm(item)
    l = fold(s)
    if not s:
        return False
    if is_admin_line(s):
        return False
    # remove TVA (sans toucher TV)
    if re.search(r"\btva\b", l) or "total tva" in l:
        return False
    if is_addressy(s) and "périmètre" not in l:
        return False
    return True


# =========================
# EXTRACTION ITEMS
# =========================
def extract_items(lines: List[str], keep_hints: List[str], max_len: int, relax: bool, tabular_labels: bool) -> List[str]:
    items: List[str] = []
    for ln in lines:
        s = norm(ln)
        if not s:
            continue

        if s.startswith(("•", "-", "–")):
            it = s.lstrip("•-– ").strip()
            if it and not is_noise_line(it):
                items.append(it)
            continue

        if looks_like_price_table_line(s) and tabular_labels:
            lab = extract_left_label_from_tabular(s)
            if lab:
                items.append(lab)
            continue

        if is_noise_line(s):
            continue

        l = fold(s)
        if len(s) <= max_len and any(k in l for k in keep_hints):
            items.append(s)
            continue

        if relax and len(s) <= max_len and sum(ch.isalpha() for ch in s) >= 10 and not is_admin_line(s):
            items.append(s)

    out, seen = [], set()
    for it in items:
        k = fold(it)
        if k in seen:
            continue
        seen.add(k)
        out.append(it)
    return out


# =========================
# ROUTING (item -> post)
# =========================
def route_catering_item(item: str) -> str:
    l = fold(item)

    if "cocktail" in l or "apéritif" in l or "aperitif" in l or "pièce" in l or "pieces" in l or "verrine" in l:
        return "Cocktail"
    if "déjeuner" in l or "dejeuner" in l or "buffet" in l or "wrap" in l or "sandwich" in l or "salade" in l or "plat" in l:
        return "Déjeuner"
    if "pause" in l or "viennoiser" in l or "gourmand" in l or "mignard" in l or "financier" in l or "cannel" in l:
        if re.search(r"\b(14h|15h|16h|17h)\b", l):
            return "Pause après-midi"
        return "Pause matin"
    if "accueil" in l or "café" in l or "cafe" in l or "thé" in l or "the" in l or "thermos" in l:
        return "Accueil café"
    if "vin" in l or "champagne" in l or "soft" in l or "jus" in l or "eau" in l:
        return "Boissons (global)"
    if "option" in l or "en option" in l or "supplément" in l or "supplement" in l:
        return "Options"
    if any(k in l for k in ["verrerie", "flûtes", "assiettes", "nappage", "mobilier", "livraison", "reprise", "personnel", "service", "mise en place"]):
        return "Autres (logistique)"
    return "Autres (logistique)"


def route_tech_item(item: str) -> str:
    l = fold(item)
    if any(k in l for k in ["réalisateur", "realisateur", "cadreur", "ingénieur", "ingenieur", "son", "technicien"]):
        return "Équipe"
    if any(k in l for k in ["caméra", "camera", "4k", "objectif", "pied", "captation"]):
        return "Captation"
    if any(k in l for k in ["régie", "regie", "mélangeur", "melangeur", "obs", "vmix", "console", "écran", "ecran", "pavlov", "zapette", "tv 55"]):
        return "Régie"
    if any(k in l for k in ["zoom", "live", "duplex", "diffusion", "plateforme", "stream"]):
        return "Diffusion"
    if any(k in l for k in ["replay", "wetransfer", "we transfer", "enregistrement", "export"]):
        return "Replay"
    if any(k in l for k in ["inclus", "comprend", "incluant"]):
        return "Inclus"
    if any(k in l for k in ["option", "supplément", "supplement", "contrainte"]):
        return "Contraintes / options"
    return "Périmètre"


# =========================
# SUMMARY GENERATION
# =========================
def extract_piece_counts(texts: List[str]) -> Dict[str, int]:
    joined = " \n ".join(texts)
    lj = fold(joined)

    def find_int(pat: str) -> Optional[int]:
        m = re.search(pat, lj, flags=re.I)
        if not m:
            return None
        try:
            return int(m.group(1))
        except Exception:
            return None

    total = find_int(r"\b(\d{1,2})\s*pi[eè]ces?\s*(par\s+personne|/pers|par\s+convive)")
    sale = find_int(r"\b(\d{1,2})\s*pi[eè]ces?.{0,30}(sal[ée]es?|froides?)")
    sucre = find_int(r"\b(\d{1,2})\s*pi[eè]ces?.{0,30}(sucr[ée]es?)")
    return {"total": total or 0, "sale": sale or 0, "sucre": sucre or 0}


def summarize_short(items: List[str], max_tokens: int = 12) -> str:
    s = " ".join(items)
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return "—"
    toks = s.split()
    if len(toks) <= max_tokens:
        return s
    return " ".join(toks[:max_tokens]) + "…"


def make_summary_catering_post(post: str, items: List[str]) -> str:
    if not items:
        return "—"

    if post == "Cocktail":
        counts = extract_piece_counts(items)
        bits = []
        if counts["total"] > 0:
            bits.append(f"{counts['total']} pièces/pers")
        if counts["sale"] > 0 or counts["sucre"] > 0:
            sub = []
            if counts["sale"] > 0:
                sub.append(f"{counts['sale']} salées")
            if counts["sucre"] > 0:
                sub.append(f"{counts['sucre']} sucrées")
            bits.append(" + ".join(sub))
        opt_sucre = any("sucr" in fold(x) and "option" in fold(x) for x in items)
        if opt_sucre:
            bits.append("(sucré en option)")
        return " — ".join(bits) if bits else "Cocktail (voir détail)"

    if post in ["Pause matin", "Pause après-midi"]:
        lj = fold(" ".join(items))
        m = re.search(r"\b(\d{1,2})\s+(gourmandis|mignardi|pi[eè]ces?)\w*\s+par\s+personne\b", lj)
        if m:
            n = m.group(1)
            ex = []
            for e in ["financier", "cannel", "cookie", "brownie", "madeleine", "mignard"]:
                if e in lj:
                    ex.append(e)
            ex_txt = f" ({', '.join(ex[:3])})" if ex else ""
            return f"{n} gourmandises/pers{ex_txt}"
        m2 = re.search(r"\b(\d{1,3})\s+financiers\b", lj)
        if m2:
            return f"{m2.group(1)} financiers"
        return summarize_short(items, 10)

    if post == "Accueil café":
        lj = fold(" ".join(items))
        flags = []
        if "café" in lj or "cafe" in lj:
            flags.append("café")
        if "thé" in lj or "the" in lj:
            flags.append("thé")
        if "jus" in lj:
            flags.append("jus")
        if "soft" in lj:
            flags.append("softs")
        return " + ".join(dict.fromkeys(flags)) if flags else summarize_short(items, 10)

    if post == "Déjeuner":
        lj = fold(" ".join(items))
        tags = []
        if "buffet" in lj:
            tags.append("buffet")
        if "wrap" in lj or "sandwich" in lj:
            tags.append("wrap/sandwich")
        if "salade" in lj:
            tags.append("salade")
        if "dessert" in lj or "tarte" in lj or "entremet" in lj:
            tags.append("dessert")
        return " + ".join(dict.fromkeys(tags)) if tags else summarize_short(items, 10)

    if post == "Boissons (global)":
        lj = fold(" ".join(items))
        tags = []
        if "soft" in lj or "jus" in lj or "eau" in lj:
            tags.append("softs/eaux")
        if "vin" in lj:
            tags.append("vin")
        if "champagne" in lj:
            tags.append("champagne (option)" if "option" in lj else "champagne")
        return " + ".join(dict.fromkeys(tags)) if tags else summarize_short(items, 10)

    if post == "Options":
        keep = []
        for it in items:
            li = fold(it)
            if "décoration" in li or "decoration" in li or "florale" in li or "suppl" in li or "option" in li:
                keep.append(it)
        return summarize_short(keep, 10) if keep else "—"

    if post == "Autres (logistique)":
        lj = fold(" ".join(items))
        tags = []
        if "livraison" in lj or "reprise" in lj:
            tags.append("livraison/reprise")
        if "personnel" in lj or "service" in lj:
            tags.append("personnel")
        if "vaisselle" in lj or "verrerie" in lj or "flûte" in lj or "assiette" in lj:
            tags.append("vaisselle/verrerie")
        if "mobilier" in lj or "mange" in lj:
            tags.append("mobilier")
        return " + ".join(dict.fromkeys(tags)) if tags else "Voir détail"

    return summarize_short(items, 10)


def make_summary_tech_post(post: str, items: List[str]) -> str:
    if not items:
        return "—"
    if post == "Équipe":
        lj = fold(" ".join(items))
        roles = []
        for r in ["réalisateur", "cadreur", "ingénieur du son", "ingenieur du son", "technicien"]:
            if r in lj:
                roles.append(r.replace("ingenieur", "ingénieur").title())
        return " + ".join(dict.fromkeys(roles)) if roles else summarize_short(items, 10)
    return summarize_short(items, 12)


# =========================
# MODEL
# =========================
@dataclass
class CateringOffer:
    vendor: str
    total_ttc: Optional[float]
    posts_detail: Dict[str, List[str]] = field(default_factory=dict)
    posts_summary: Dict[str, str] = field(default_factory=dict)
    comment: str = ""


@dataclass
class TechOffer:
    vendor: str
    total_ttc: Optional[float]
    posts_detail: Dict[str, List[str]] = field(default_factory=dict)
    posts_summary: Dict[str, str] = field(default_factory=dict)
    comment: str = ""


def parse_catering_offer(text: str, filename: str) -> CateringOffer:
    vendor = guess_vendor_name(text, filename)
    if vendor_is_suspicious(vendor):
        vendor = vendor_from_filename(filename)

    total_ttc = find_total_ttc(text)

    lines = cut_at_cgv(split_lines(text))
    content = [ln for ln in lines if not is_noise_line(ln)]

    raw_items = extract_items(content, MENU_HINTS, max_len=320, relax=True, tabular_labels=True)

    posts_detail = {p: [] for p in CATERING_POSTS}
    for it in raw_items:
        if not should_keep_catering_detail(it):
            continue
        posts_detail[route_catering_item(it)].append(it)

    for p in posts_detail:
        seen, out = set(), []
        for it in posts_detail[p]:
            k = fold(it)
            if k in seen:
                continue
            seen.add(k)
            out.append(it)
        posts_detail[p] = out

    posts_summary = {p: make_summary_catering_post(p, posts_detail[p]) for p in CATERING_POSTS}
    return CateringOffer(vendor=vendor, total_ttc=total_ttc, posts_detail=posts_detail, posts_summary=posts_summary)


def parse_tech_offer(text: str, filename: str) -> TechOffer:
    vendor = guess_vendor_name(text, filename)
    if vendor_is_suspicious(vendor):
        vendor = vendor_from_filename(filename)

    total_ttc = find_total_ttc(text)

    lines = cut_at_cgv(split_lines(text))
    content = [ln for ln in lines if not is_noise_line(ln)]

    raw_items = extract_items(content, TECH_HINTS, max_len=420, relax=False, tabular_labels=True)

    posts_detail = {p: [] for p in TECH_POSTS}
    for it in raw_items:
        if not should_keep_tech_detail(it):
            continue
        posts_detail[route_tech_item(it)].append(it)

    for p in posts_detail:
        seen, out = set(), []
        for it in posts_detail[p]:
            k = fold(it)
            if k in seen:
                continue
            seen.add(k)
            out.append(it)
        posts_detail[p] = out

    posts_summary = {p: make_summary_tech_post(p, posts_detail[p]) for p in TECH_POSTS}
    return TechOffer(vendor=vendor, total_ttc=total_ttc, posts_detail=posts_detail, posts_summary=posts_summary)


# =========================
# WORD HELPERS
# =========================
def set_cell_shading(cell, fill_hex: str):
    fill_hex = fill_hex.replace("#", "")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_run(run, bold=False, size=10, color="#111827"):
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = hex_to_rgb(color)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=16, color=PRIMARY)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=11, color="#111827")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_small(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=False, size=9, color="#374151")


def add_band(doc: Document, label: str, sub: str = ""):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, PRIMARY)
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    set_run(r, bold=True, size=11, color="#FFFFFF")
    if sub:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(sub)
        set_run(r2, bold=False, size=9, color="#FFFFFF")
    doc.add_paragraph("")


def add_offer_detail_table(doc: Document, posts: Dict[str, List[str]], order: List[str]):
    t = doc.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = "Poste"
    h[1].text = "Détail"
    set_cell_shading(h[0], PRIMARY)
    set_cell_shading(h[1], PRIMARY)
    for cell in h:
        for rr in cell.paragraphs[0].runs:
            set_run(rr, bold=True, size=10, color="#FFFFFF")
        set_cell_margins(cell)

    for post in order:
        rr = t.add_row().cells
        rr[0].text = post
        set_cell_shading(rr[0], "F3F4F6")
        for run in rr[0].paragraphs[0].runs:
            set_run(run, bold=True, size=9, color="#111827")
        set_cell_margins(rr[0])

        items = posts.get(post, [])
        rr[1].text = "—" if not items else "\n".join(f"• {x}" for x in items)
        for p in rr[1].paragraphs:
            for run in p.runs:
                set_run(run, bold=False, size=9, color="#111827")
        set_cell_margins(rr[1])

    doc.add_paragraph("")


def build_word(event_title: str, event_date: str, guests: int,
               catering: List[CateringOffer], tech: List[TechOffer]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    add_title(doc, "SYNTHÈSE DEVIS — PRESTATAIRES")
    add_subtitle(doc, f"{event_title} — {event_date} — Sur la base de {guests} convives")
    add_small(doc, f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    # Traiteur: SUMMARY
    if catering:
        add_subtitle(doc, "1) PRESTATION TRAITEUR — Comparatif (synthèse)")
        vendors = catering[:3]
        table = doc.add_table(rows=1, cols=1 + len(vendors))

        hdr = table.rows[0].cells
        hdr[0].text = "Poste"
        set_cell_shading(hdr[0], PRIMARY)
        for rr in hdr[0].paragraphs[0].runs:
            set_run(rr, bold=True, size=10, color="#FFFFFF")

        for i, off in enumerate(vendors, start=1):
            hdr[i].text = off.vendor
            set_cell_shading(hdr[i], PRIMARY)
            for rr in hdr[i].paragraphs[0].runs:
                set_run(rr, bold=True, size=10, color="#FFFFFF")
        for c in hdr:
            set_cell_margins(c)

        rows = ["Total TTC"] + CATERING_POSTS + ["Commentaire"]
        for label in rows:
            r = table.add_row().cells
            r[0].text = label
            set_cell_shading(r[0], "F3F4F6")
            for rr in r[0].paragraphs[0].runs:
                set_run(rr, bold=True, size=9, color="#111827")
            set_cell_margins(r[0])

            for j, off in enumerate(vendors, start=1):
                if label == "Total TTC":
                    val = euro_fmt(off.total_ttc)
                elif label == "Commentaire":
                    val = norm(off.comment) or "—"
                else:
                    val = off.posts_summary.get(label, "—") or "—"
                r[j].text = val
                for p in r[j].paragraphs:
                    for rr in p.runs:
                        set_run(rr, bold=False, size=9, color="#111827")
                set_cell_margins(r[j])

        doc.add_paragraph("")

    # Technique: SUMMARY
    if tech:
        add_subtitle(doc, "2) PRESTATION TECHNIQUE — Synthèse")
        for idx, off in enumerate(tech[:2], start=1):
            p = doc.add_paragraph()
            rr = p.add_run(f"Prestataire technique {idx} : {off.vendor} — Total TTC : {euro_fmt(off.total_ttc)}")
            set_run(rr, bold=True, size=9, color="#111827")

            t = doc.add_table(rows=1, cols=2)
            h = t.rows[0].cells
            h[0].text = "Item"
            h[1].text = "Synthèse"
            set_cell_shading(h[0], PRIMARY)
            set_cell_shading(h[1], PRIMARY)
            for cell in h:
                for rrr in cell.paragraphs[0].runs:
                    set_run(rrr, bold=True, size=10, color="#FFFFFF")
                set_cell_margins(cell)

            for item in TECH_POSTS:
                rrrow = t.add_row().cells
                rrrow[0].text = item
                set_cell_shading(rrrow[0], "F3F4F6")
                for rrr in rrrow[0].paragraphs[0].runs:
                    set_run(rrr, bold=True, size=9, color="#111827")
                set_cell_margins(rrrow[0])

                val = norm(off.comment) if item == "Conseil" else off.posts_summary.get(item, "—")
                rrrow[1].text = val if val else "—"
                set_cell_margins(rrrow[1])
                for p in rrrow[1].paragraphs:
                    for rrr in p.runs:
                        set_run(rrr, bold=False, size=9, color="#111827")
            doc.add_paragraph("")

    # DETAILS (IMPORTANT: plus de saut de page juste après l'intro)
    doc.add_page_break()
    add_title(doc, "DÉTAIL DES OFFRES (modifiable via l’outil)")
    add_small(doc, "Le comparatif ci-dessus est une synthèse ; ci-dessous la liste des contenus.")
    doc.add_paragraph("")

    if catering:
        # PAS de doc.add_page_break() ici => plus de page blanche
        add_title(doc, "DÉTAIL — PRESTATIONS TRAITEUR")
        add_small(doc, "Listes complètes (pièces cocktail, menus, art de la table, logistique).")
        doc.add_paragraph("")
        for off in catering[:3]:
            add_band(doc, off.vendor, f"Total TTC : {euro_fmt(off.total_ttc)}")
            add_offer_detail_table(doc, off.posts_detail, CATERING_POSTS)

    if tech:
        doc.add_page_break()
        add_title(doc, "DÉTAIL — PRESTATIONS TECHNIQUES")
        add_small(doc, "Listes complètes.")
        doc.add_paragraph("")
        for off in tech[:2]:
            add_band(doc, off.vendor, f"Total TTC : {euro_fmt(off.total_ttc)}")
            add_offer_detail_table(doc, off.posts_detail, TECH_POSTS)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================
# STREAMLIT
# =========================
st.set_page_config(page_title=APP_TITLE, layout="wide")

st.markdown(
    f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700;800;900&display=swap');
html, body, .stApp, [class*="css"] {{ font-family: 'Montserrat', sans-serif !important; }}
header[data-testid="stHeader"] {{ display: none; }}
.stApp {{ background: {BG}; }}
[data-testid="stHorizontalBlock"] {{
  background: white; border-radius: 16px; padding: 0.75rem 0.85rem;
  margin-bottom: 0.65rem; box-shadow: 0 1px 12px rgba(0,0,0,0.06);
}}
.stButton > button {{
  background-color: {PRIMARY} !important; color: #ffffff !important; border: none !important;
  border-radius: 14px !important; padding: 0.85rem 1.05rem !important; font-weight: 900 !important;
  min-height: 52px !important; white-space: nowrap !important;
}}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(f"## {APP_TITLE}")
st.caption("Comparatif = résumé. Détail = listes complètes. (Intro détails sans page blanche.)")
st.divider()

ttc_min = st.number_input("Seuil TTC minimum (alerte) — pas bloquant", min_value=0, max_value=200000, value=500, step=50)

c1, c2, c3 = st.columns([2, 1, 1], vertical_alignment="center")
with c1:
    event_title = st.text_input("Événement (titre court)", placeholder="Ex : Journée scientifique — Colloque Génétique et Société")
with c2:
    event_date = st.text_input("Date", placeholder="Ex : 19/03/2026")
with c3:
    guests = st.number_input("Nb convives", min_value=1, max_value=5000, value=100, step=10)

st.markdown("### 1) Devis traiteur (max 3)")
catering_files = st.file_uploader("Upload PDF traiteur", type=["pdf"], accept_multiple_files=True)

st.markdown("### 2) Devis technique (max 2)")
tech_files = st.file_uploader("Upload PDF technique", type=["pdf"], accept_multiple_files=True)

if (not catering_files) and (not tech_files):
    st.info("Commence par uploader au moins un devis PDF.")
    st.stop()

catering_files = (catering_files or [])[:3]
tech_files = (tech_files or [])[:2]

with st.spinner("Lecture des PDFs…"):
    catering_offers: List[CateringOffer] = []
    tech_offers: List[TechOffer] = []

    for f in catering_files:
        txt = extract_pdf_text(f)
        catering_offers.append(parse_catering_offer(txt, f.name))

    for f in tech_files:
        txt = extract_pdf_text(f)
        tech_offers.append(parse_tech_offer(txt, f.name))

tabA, tabB, tabC = st.tabs(["Traiteur — Synthèse", "Traiteur — Détail", "Technique"])

with tabA:
    if not catering_offers:
        st.caption("Aucun devis traiteur.")
    else:
        for i, off in enumerate(catering_offers, start=1):
            with st.expander(f"Traiteur {i} — {off.vendor}", expanded=(i == 1)):
                off.vendor = st.text_input("Nom prestataire", value=off.vendor, key=f"c_vendor_{i}")
                if vendor_is_suspicious(off.vendor):
                    st.warning("Nom prestataire suspect. (Astuce: renomme le PDF avec le nom du prestataire.)")

                ttc_in = st.text_input("Total TTC", value=("" if off.total_ttc is None else euro_fmt(off.total_ttc)), key=f"c_ttc_{i}")
                off.total_ttc = parse_eur_amount(ttc_in)
                if off.total_ttc is not None and off.total_ttc < float(ttc_min):
                    st.warning(f"TTC < {ttc_min}€ : probable mauvaise détection.")

                st.markdown("**Synthèse par poste (alimentation du 1er tableau)**")
                for post in CATERING_POSTS:
                    off.posts_summary[post] = st.text_input(
                        f"{post} — Synthèse",
                        value=off.posts_summary.get(post, "—"),
                        key=f"c_sum_{i}_{post}",
                    )
                off.comment = st.text_area("Commentaire", value=off.comment, height=80, key=f"c_comment_{i}")

with tabB:
    if not catering_offers:
        st.caption("Aucun devis traiteur.")
    else:
        st.info("Ici tu modifies les listes complètes affichées dans la section DÉTAIL du Word.")
        for i, off in enumerate(catering_offers, start=1):
            with st.expander(f"Détail — Traiteur {i} : {off.vendor}", expanded=False):
                for post in CATERING_POSTS:
                    edited = st.text_area(
                        f"{post} — Détail (liste)",
                        value="\n".join(off.posts_detail.get(post, [])),
                        height=160,
                        key=f"c_det_{i}_{post}",
                    )
                    off.posts_detail[post] = [norm(x) for x in edited.splitlines() if norm(x)]

with tabC:
    if not tech_offers:
        st.caption("Aucun devis technique.")
    else:
        for i, off in enumerate(tech_offers, start=1):
            with st.expander(f"Technique {i} — {off.vendor}", expanded=(i == 1)):
                off.vendor = st.text_input("Nom prestataire", value=off.vendor, key=f"t_vendor_{i}")
                if vendor_is_suspicious(off.vendor):
                    st.warning("Nom prestataire suspect. (Astuce: renomme le PDF.)")

                ttc_in = st.text_input("Total TTC", value=("" if off.total_ttc is None else euro_fmt(off.total_ttc)), key=f"t_ttc_{i}")
                off.total_ttc = parse_eur_amount(ttc_in)

                st.markdown("**Synthèse technique**")
                for post in TECH_POSTS:
                    off.posts_summary[post] = st.text_input(
                        f"{post} — Synthèse",
                        value=off.posts_summary.get(post, "—"),
                        key=f"t_sum_{i}_{post}",
                    )

                st.markdown("**Détail technique (listes)**")
                for post in TECH_POSTS:
                    edited = st.text_area(
                        f"{post} — Détail (liste)",
                        value="\n".join(off.posts_detail.get(post, [])),
                        height=130,
                        key=f"t_det_{i}_{post}",
                    )
                    off.posts_detail[post] = [norm(x) for x in edited.splitlines() if norm(x)]

                off.comment = st.text_area("Conseil", value=off.comment, height=90, key=f"t_comment_{i}")

st.divider()
if st.button("Générer le Word (.docx)", use_container_width=True, type="primary"):
    docx_bytes = build_word(
        event_title=event_title.strip() or "Événement",
        event_date=event_date.strip() or "Date à préciser",
        guests=int(guests),
        catering=catering_offers,
        tech=tech_offers,
    )
    ts = datetime.now().strftime("%Y-%m-%d_%H%M")
    st.download_button(
        "⬇️ Télécharger la synthèse (Word)",
        data=docx_bytes,
        file_name=f"synthese_devis_{ts}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
